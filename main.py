from flask import Flask, render_template, request, send_file
import speech_recognition as sr
from datetime import datetime
from docx import Document
from fpdf import FPDF
from io import BytesIO
import re
import wave

app = Flask(__name__, template_folder="templates", static_folder="static")

# Keywords to identify agenda points
KEYWORDS = ["discussion", "decision", "action", "task", "update", "review", "agenda", "issue"]

# Function to transcribe WAV audio using SpeechRecognition
def transcribe_audio(wav_file):
    recognizer = sr.Recognizer()
    audio_file = sr.AudioFile(wav_file)
    with audio_file as source:
        audio_data = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio_data)
    except sr.UnknownValueError:
        text = ""
    return text

# Function to format transcription into agenda items
def format_agenda(transcription):
    lines = re.split(r'[.?!]\s+', transcription)
    agenda = []
    current_item = []
    item_count = 1
    for line in lines:
        if any(word.lower() in line.lower() for word in KEYWORDS):
            if current_item:
                agenda.append(f"Agenda Item {item_count}:\n- " + "\n- ".join(current_item))
                item_count += 1
                current_item = []
        current_item.append(line.strip())
    if current_item:
        agenda.append(f"Agenda Item {item_count}:\n- " + "\n- ".join(current_item))
    return "\n\n".join(agenda)


# Function to save agenda as Word
def save_as_word(text, meeting_date, meeting_time):
    doc = Document()
    doc.add_heading("Meeting Minutes", 0)
    doc.add_paragraph(f"Date: {meeting_date}")
    doc.add_paragraph(f"Time: {meeting_time}")
    doc.add_paragraph("")  # Empty line
    for item in text.split("\n\n"):
        heading, *points = item.split("\n- ")
        doc.add_heading(heading, level=1)
        for point in points:
            if point.strip():
                doc.add_paragraph(point.strip(), style="List Bullet")
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# Function to save agenda as PDF
def save_as_pdf(text, meeting_date, meeting_time):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)  # TrueType Unicode font
    pdf.set_font("DejaVu", '', 16)
    pdf.multi_cell(0, 10, "Meeting Minutes\n\n")
    pdf.set_font("DejaVu", '', 12)
    pdf.multi_cell(0, 10, f"Date: {meeting_date}")
    pdf.multi_cell(0, 10, f"Time: {meeting_time}\n")

    for item in text.split("\n\n"):
        heading, *points = item.split("\n- ")
        pdf.set_font("DejaVu", '', 14)
        pdf.multi_cell(0, 10, heading)
        pdf.set_font("DejaVu", '', 12)
        for point in points:
            if point.strip():
                pdf.multi_cell(0, 10, f"• {point.strip()}")
        pdf.ln(5)

    pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
    file_stream = BytesIO(pdf_bytes)
    file_stream.seek(0)
    return file_stream
    return file_stream

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        file = request.files["audio_file"]
        format_type = request.form.get("format")
        
        # Only accept WAV files
        if not file.filename.lower().endswith(".wav"):
            return "Error: Only WAV files are supported.", 400
        
        # Optional: validate WAV file can be read
        try:
            with wave.open(file, 'rb') as wav:
                pass
        except wave.Error:
            return "Error: Invalid WAV file.", 400
        
        # Reset file pointer for speech recognition
        file.seek(0)
        transcription = transcribe_audio(file)
        agenda_text = format_agenda(transcription)

        now = datetime.now()
        meeting_date = request.form.get("meeting_date", now.strftime("%d-%m-%Y"))
        meeting_time = request.form.get("meeting_time", now.strftime("%H:%M %p"))
        
        if format_type == "pdf":
            file_stream = save_as_pdf(agenda_text, meeting_date, meeting_time)
            return send_file(file_stream, as_attachment=True, download_name="agenda.pdf")
        else:
            file_stream = save_as_word(agenda_text, meeting_date, meeting_time)
            return send_file(file_stream, as_attachment=True, download_name="agenda.docx")
        
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
